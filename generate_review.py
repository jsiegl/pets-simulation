"""
Generates PET Simulation Review.docx in the same folder as this script.
Run: python generate_review.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ── Page margins ───────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1.2)
    section.right_margin  = Inches(1.2)

# ── Colour palette ─────────────────────────────────────────────────────────────
DARK_BLUE  = RGBColor(0x0A, 0x12, 0x28)
MID_BLUE   = RGBColor(0x1E, 0x2F, 0x50)
ACCENT     = RGBColor(0x00, 0xB5, 0x80)   # teal / mint
CORAL      = RGBColor(0xC0, 0x39, 0x50)
AMBER      = RGBColor(0xB8, 0x8A, 0x00)
BLUE_TEXT  = RGBColor(0x1F, 0x62, 0xB5)
GREY_TEXT  = RGBColor(0x44, 0x44, 0x44)
BLACK      = RGBColor(0x11, 0x11, 0x11)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_GREY = RGBColor(0xF2, 0xF4, 0xF8)
ROW_ALT    = RGBColor(0xEB, 0xF0, 0xFA)


# ── Helper: set paragraph shading ─────────────────────────────────────────────
def shade_paragraph(para, fill_rgb: RGBColor):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    hex_color = '{:02X}{:02X}{:02X}'.format(fill_rgb[0], fill_rgb[1], fill_rgb[2])
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    pPr.append(shd)


# ── Helper: shade a table cell ─────────────────────────────────────────────────
def shade_cell(cell, fill_rgb: RGBColor):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    hex_color = '{:02X}{:02X}{:02X}'.format(fill_rgb[0], fill_rgb[1], fill_rgb[2])
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)


# ── Helper: set table cell borders ────────────────────────────────────────────
def set_cell_border(cell, **kwargs):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        if edge in kwargs:
            tag = OxmlElement(f'w:{edge}')
            for attr, val in kwargs[edge].items():
                tag.set(qn(f'w:{attr}'), str(val))
            tcBorders.append(tag)
    tcPr.append(tcBorders)


def thin_border(cell):
    border = {'val': 'single', 'sz': '4', 'space': '0', 'color': '8EA8C3'}
    set_cell_border(cell, top=border, bottom=border, left=border, right=border)


# ── Helper: run with colour ────────────────────────────────────────────────────
def add_run(para, text, bold=False, italic=False, color=None, size=None, mono=False):
    run = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    if mono:
        run.font.name = 'Courier New'
    return run


# ── Helper: styled heading ─────────────────────────────────────────────────────
def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(4)
    shade_paragraph(p, DARK_BLUE)
    run = p.add_run(f'  {text}')
    run.bold = True
    run.font.size  = Pt(15)
    run.font.color.rgb = ACCENT
    run.font.name  = 'Calibri'
    return p


def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.bold = True
    run.font.size  = Pt(12)
    run.font.color.rgb = DARK_BLUE
    run.font.name  = 'Calibri'
    # bottom border
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '6')
    bot.set(qn('w:space'), '4')
    bot.set(qn('w:color'), '1E2F50')
    pBdr.append(bot)
    pPr.append(pBdr)
    return p


def h3(text, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.font.size  = Pt(10.5)
    run.font.color.rgb = color or BLUE_TEXT
    run.font.name  = 'Calibri'
    return p


def body(text, left_indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    if left_indent:
        p.paragraph_format.left_indent = Inches(0.25)
    run = p.add_run(text)
    run.font.size  = Pt(10)
    run.font.color.rgb = GREY_TEXT
    run.font.name  = 'Calibri'
    return p


def bullet(text, indent=0, color=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Inches(0.3 + indent * 0.2)
    run = p.add_run(text)
    run.font.size  = Pt(10)
    run.font.color.rgb = color or GREY_TEXT
    run.font.name  = 'Calibri'
    return p


def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Inches(0.25)
    shade_paragraph(p, RGBColor(0xEE, 0xF1, 0xF8))
    run = p.add_run(text)
    run.font.name  = 'Courier New'
    run.font.size  = Pt(8.5)
    run.font.color.rgb = RGBColor(0x0A, 0x30, 0x60)
    return p


def callout(label, text, fill=None, label_color=None):
    fill = fill or RGBColor(0xE8, 0xF4, 0xFD)
    label_color = label_color or BLUE_TEXT
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.left_indent  = Inches(0.3)
    shade_paragraph(p, fill)
    r1 = p.add_run(f'{label}  ')
    r1.bold = True
    r1.font.size  = Pt(10)
    r1.font.color.rgb = label_color
    r1.font.name  = 'Calibri'
    r2 = p.add_run(text)
    r2.font.size  = Pt(10)
    r2.font.color.rgb = GREY_TEXT
    r2.font.name  = 'Calibri'


def spacer():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)


# ══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(36)
p.paragraph_format.space_after  = Pt(6)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
shade_paragraph(p, DARK_BLUE)
r = p.add_run('  Privacy Enhancing Technologies  ')
r.bold = True
r.font.size  = Pt(22)
r.font.color.rgb = WHITE
r.font.name  = 'Calibri'

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
shade_paragraph(p2, DARK_BLUE)
r2 = p2.add_run('  Interactive Simulation — Code & Content Review  ')
r2.bold = True
r2.font.size  = Pt(14)
r2.font.color.rgb = ACCENT
r2.font.name  = 'Calibri'

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
shade_paragraph(p3, DARK_BLUE)
r3 = p3.add_run('  Claude Code Review  ·  April 2026  ')
r3.font.size  = Pt(10)
r3.font.color.rgb = RGBColor(0x88, 0xA0, 0xC8)
r3.font.name  = 'Calibri'

spacer()

p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = p4.add_run('Reviewed files: index.html · app.js · content.js · styles.css')
r4.italic = True
r4.font.size  = Pt(9)
r4.font.color.rgb = RGBColor(0x77, 0x88, 0xAA)
r4.font.name  = 'Calibri'

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — SCIENTIFIC ACCURACY
# ══════════════════════════════════════════════════════════════════════════════

h1('1  Scientific Accuracy')

body('Analysis of whether the simulation logic accurately reflects real-world PET principles.')

# ── 1.1 Global DP ──────────────────────────────────────────────────────────────
h2('1.1  Differential Privacy — Global (Laplace Mechanism)')

body('The inverse-CDF Laplace sampler implementation is mathematically correct:')

code_block(
    'const scale = 1 / epsilon;\n'
    'const u = Math.random() - 0.5;\n'
    'const noise = -scale * Math.sign(u) * Math.log(1 - 2 * Math.abs(u));\n'
    'return Math.max(0, Math.round(v + noise));   // app.js:472–478'
)

body('However, three inaccuracies exist:')

bullet('Sensitivity is silently fixed at 1. The Laplace scale is b = sensitivity / ε, but sensitivity = 1 is assumed without explanation or UI control. For histogram queries over student counts in the hundreds, the choice of global sensitivity 1 is not self-evident and should be surfaced.')
bullet('Math.max(0, …) clipping introduces positive bias. Truncating negative noisy counts to zero is a post-processing step that biases results upward for small groups — a well-known DP pitfall that is not flagged anywhere in the UI.')
bullet('No privacy budget composition tracking. Each "Run Query" click draws fresh Laplace noise, but the UI never accumulates the total ε spent. Under sequential composition, running the query k times costs k × ε. A user clicking "Run" ten times at ε = 1.0 has actually spent ε = 10.0.')

# ── 1.2 Local DP Bug ──────────────────────────────────────────────────────────
h2('1.2  Differential Privacy — Local (Randomized Response)  ⚠ Bug Found')

callout('Bug:', 'The Local DP correction formula uses the wrong denominator, producing a biased corrected-rate estimate.',
        fill=RGBColor(0xFD, 0xEE, 0xEE), label_color=CORAL)

body('The code at app.js:673 implements:')
code_block('const correctedRate = (observedRate - (1 - p) / 2) / (2 * p - 1);   // INCORRECT')

body('For the described mechanism (honest coin then random second flip), the mathematically correct derivation is:')
code_block(
    'P(report=1) = p × π + (1−p) × 0.5\n'
    '⟹  π  =  (λ − (1−p)/2) / p           // correct denominator: p\n\n'
    '// Code uses (2p−1) as denominator — wrong for this mechanism'
)

body('Impact for p = 0.7 (the configured value):')

tbl = doc.add_table(rows=4, cols=3)
tbl.style = 'Table Grid'
tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

headers = ['', 'Formula', 'Result for true rate = 35%']
for ci, h in enumerate(headers):
    cell = tbl.cell(0, ci)
    shade_cell(cell, DARK_BLUE)
    thin_border(cell)
    run = cell.paragraphs[0].add_run(h)
    run.bold = True
    run.font.color.rgb = WHITE
    run.font.size = Pt(9)
    run.font.name = 'Calibri'

rows_data = [
    ('Correct formula',       'denominator = p = 0.70',               '≈ 35.0%  ✓'),
    ('Code (app.js:673)',     'denominator = 2p−1 = 0.40',            '≈ 61.3%  ✗'),
    ('content.js comment',   'denominator = p−(1−p)/2 = 0.55',       '≈ 44.5%  ✗'),
]
fills = [LIGHT_GREY, RGBColor(0xFD, 0xF0, 0xF0), RGBColor(0xFD, 0xF0, 0xF0)]
for ri, (label, formula, result) in enumerate(rows_data):
    for ci, text in enumerate([label, formula, result]):
        cell = tbl.cell(ri + 1, ci)
        shade_cell(cell, fills[ri])
        thin_border(cell)
        run = cell.paragraphs[0].add_run(text)
        run.font.size = Pt(9)
        run.font.name = 'Courier New' if ci == 1 else 'Calibri'
        if ci == 2 and ri > 0:
            run.font.color.rgb = CORAL
        elif ci == 2:
            run.font.color.rgb = ACCENT

spacer()
body('Additionally, content.js documents a third formula in comments that differs from both the correct formula and the implementation — three inconsistent values across two files, with only one being correct.')

# ── 1.3 MPC ──────────────────────────────────────────────────────────────────
h2('1.3  Secure MPC — Secret Sharing')

body('The share generation code comment reads "// Show shares (fake split values)" — the visualization is cosmetic. Share chips display plausible-looking splits (S1→P1: 20, S1→P2: 15…) but:')
bullet('No additive verification is ever performed in the UI. The user cannot confirm shares reconstruct the correct sum.')
bullet('Individual school scores are correctly hidden in the UI, but the protocol outcome (district average) is computed outside any simulated MPC protocol — it is just JavaScript arithmetic.')
bullet('In real additive secret sharing, shares can be negative. This is not conveyed, leaving users with an incomplete mental model.')

# ── 1.4 PPRL ─────────────────────────────────────────────────────────────────
h2('1.4  PPRL — Bloom Filter Encoding')

body('Two accuracy gaps exist in the PPRL module:')

code_block(
    '// app.js:814 — matching is done by integer id comparison, NOT by hash\n'
    'const matchedIds = agencyA.filter(a => agencyB.some(b => b.id === a.id));'
)

bullet('The fakeHash() function generates a deterministic hex string displayed in the UI, but actual record matching is performed by comparing integer id fields directly. The hash visual is purely cosmetic — it misleads users into believing Bloom filter comparison drives the match.')
bullet('Real Bloom filter PPRL supports fuzzy/approximate matching (name typos, DOB variations). The demo implies hashes must collide exactly, losing the key educational point about approximate string matching across agency boundaries.')

# ── 1.5 HE ───────────────────────────────────────────────────────────────────
h2('1.5  Homomorphic Encryption — Deterministic "Ciphertext"')

body('The fakeEnc() function generates a deterministic string from each plaintext value. Real HE schemes (BFV, CKKS, BGV) are probabilistically encrypted — the same plaintext produces a different ciphertext on every encryption call (semantic/IND-CPA security). The simulation inadvertently demonstrates semantically insecure encryption, which is the opposite of the HE property being taught.')

# ── 1.6 Synthetic Data ────────────────────────────────────────────────────────
h2('1.6  Synthetic Data — Hardcoded Perfect Fidelity')

body('The statistical comparison panel hardcodes perfectly identical figures (Real: 3.05 · Synthetic: 3.05) for every metric. Real synthetic generation involves fidelity trade-offs that worsen for small subgroups. No parameter controls synthesis quality, so the privacy-utility tension — the central lesson — is completely absent from this module.')


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — CODE EFFICIENCY
# ══════════════════════════════════════════════════════════════════════════════

h1('2  Code Efficiency')

body('Identified bottlenecks and structural issues in the simulation engine.')

tbl2 = doc.add_table(rows=8, cols=3)
tbl2.style = 'Table Grid'
tbl2.alignment = WD_TABLE_ALIGNMENT.LEFT

col_widths = [Inches(2.4), Inches(2.4), Inches(1.4)]

hdrs2 = ['Issue', 'Location', 'Impact']
for ci, h in enumerate(hdrs2):
    cell = tbl2.cell(0, ci)
    shade_cell(cell, MID_BLUE)
    thin_border(cell)
    run = cell.paragraphs[0].add_run(h)
    run.bold = True
    run.font.color.rgb = WHITE
    run.font.size = Pt(9.5)
    run.font.name = 'Calibri'

rows2 = [
    ('Full module innerHTML rebuild on every tab/module switch',
     'showModule() + all render*() functions',
     'High — discards all DOM state; no incremental update'),
    ('Global window.* function pollution (~15 functions)',
     'window.dpRunQuery, window.pprlLink, window.flRun…',
     'Medium — naming collision risk; stale functions survive module switch'),
    ('Full D3 bundle (~350 KB) loaded for only 2 modules',
     'index.html:107 — d3.min.js',
     'High — DP bar chart and FL diagram are the only D3 consumers'),
    ('No debounce on epsilon slider',
     'app.js:441–445',
     'Low now — dpDrawBars fires synchronously on every pixel of drag'),
    ('Async sleep() chains are not cancellable',
     'All async simulation functions',
     'Medium — module switch mid-animation leaves orphaned DOM mutations'),
    ('MPC chip ID sanitization is fragile',
     'app.js:905 — .replace(/[^a-z0-9]/gi,\'\')',
     'Low — labels differing only in punctuation would collide on getElementById'),
    ('Massive inline style strings throughout render functions',
     'All render* functions',
     'Medium — prevents CSS deduplication; inflates JS payload'),
]

alt_fill = [LIGHT_GREY, ROW_ALT]
for ri, (issue, loc, impact) in enumerate(rows2):
    fill = alt_fill[ri % 2]
    for ci, text in enumerate([issue, loc, impact]):
        cell = tbl2.cell(ri + 1, ci)
        shade_cell(cell, fill)
        thin_border(cell)
        run = cell.paragraphs[0].add_run(text)
        run.font.size = Pt(9)
        run.font.name = 'Courier New' if ci == 1 else 'Calibri'
        run.font.color.rgb = GREY_TEXT

spacer()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — VISUALIZATION & UX
# ══════════════════════════════════════════════════════════════════════════════

h1('3  Visualization & UX — Privacy vs. Utility Trade-Off')

h2('3.1  What Works Well')
bullet('The DP epsilon slider + side-by-side bar chart is the strongest interactive element — noise visually expands as ε drops, making the trade-off visceral and immediately understandable.')
bullet('Light/dark mode toggle with localStorage persistence is polished.')
bullet('The keyboard-accessible tooltip system (hover preview + click-to-pin) is well implemented.')
bullet('Step-by-step animation sequencing is pedagogically clear across all modules.')

h2('3.2  Gaps and Weaknesses')

callout('Critical:', '"Output type" shows *TBD* in every one of the 8 module tradeoff rows. This is visible in the live UI and undermines credibility with the target audience (state education agency staff).',
        fill=RGBColor(0xFD, 0xEE, 0xEE), label_color=CORAL)

spacer()
h3('No interactive privacy-utility control outside of DP')
body('The ε slider is the only control that lets users feel the trade-off. All other modules have binary run/reset interactions. Suggested additions:')
bullet('Synthetic Data: a fidelity slider (fewer synthetic records → lower privacy risk, lower utility) with live statistical drift visible in the comparison panel.')
bullet('k-Anonymity: the k-slider updates badge counts but does not recalculate data utility loss (% of records suppressed or further generalized).')
bullet('Federated Learning: a DP noise multiplier on gradients showing accuracy vs. ε at the model level.')

h3('True rate visible to "collector" in Local DP demo')
body('The corrected rate is displayed side-by-side with "True rate (hidden from collector)" — which is nevertheless shown in the UI. In a real system, the true rate is never observable by the collector. Displaying it frames LDP as an approximation trick rather than a genuine privacy guarantee.')

h3('No cumulative privacy budget meter')
body('When a user runs the DP query multiple times, a running ε-expenditure bar (e.g., "Budget used: 3.0 / 5.0 ε") would teach sequential composition without requiring a written explanation.')

h3('Comparison table has no interactive sorting or filtering')
body('The Compare module is a static table. For the target audience (analysts choosing between techniques), a radar chart or parallel coordinates plot would allow direct comparison on user-prioritised dimensions — e.g., dragging "Implementation complexity" to the top axis for a resource-constrained agency.')

h2('3.3  Suggested Enhancements')

tbl3 = doc.add_table(rows=6, cols=2)
tbl3.style = 'Table Grid'
tbl3.alignment = WD_TABLE_ALIGNMENT.LEFT

for ci, h in enumerate(['Enhancement', 'Educational Value']):
    cell = tbl3.cell(0, ci)
    shade_cell(cell, MID_BLUE)
    thin_border(cell)
    run = cell.paragraphs[0].add_run(h)
    run.bold = True
    run.font.color.rgb = WHITE
    run.font.size = Pt(9.5)
    run.font.name = 'Calibri'

enhancements = [
    ('Replace Synthetic Data static stats with a live chart that shifts as a fidelity parameter changes',
     'Directly demonstrates privacy-utility Pareto trade-off'),
    ('Add a "What does an adversary see?" toggle on PPRL — switch perspective to an interceptor who only receives hash blobs',
     'Teaches threat-model thinking; clarifies what "crossing the boundary" means'),
    ('Add a gradient-leak panel in FL showing that Δw can approximate training data without DP',
     'Motivates DP-on-gradients; closes a major conceptual gap'),
    ('Add a cumulative ε budget meter to the DP module',
     'Teaches composition theorem interactively'),
    ('Radar / spider chart in Compare replacing the static table',
     'Allows analysts to weight dimensions by their own priorities'),
]
for ri, (enh, val) in enumerate(enhancements):
    fill = alt_fill[ri % 2]
    for ci, text in enumerate([enh, val]):
        cell = tbl3.cell(ri + 1, ci)
        shade_cell(cell, fill)
        thin_border(cell)
        run = cell.paragraphs[0].add_run(text)
        run.font.size = Pt(9)
        run.font.name = 'Calibri'
        run.font.color.rgb = GREY_TEXT

spacer()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 4 — GAP ANALYSIS
# ══════════════════════════════════════════════════════════════════════════════

h1('4  Gap Analysis — Missing PET Concepts & Attack Vectors')

h2('4.1  Attack Vectors Not Demonstrated')

tbl4 = doc.add_table(rows=7, cols=3)
tbl4.style = 'Table Grid'
tbl4.alignment = WD_TABLE_ALIGNMENT.LEFT

for ci, h in enumerate(['Missing Attack', 'Relevant Module(s)', 'Why It Matters']):
    cell = tbl4.cell(0, ci)
    shade_cell(cell, DARK_BLUE)
    thin_border(cell)
    run = cell.paragraphs[0].add_run(h)
    run.bold = True
    run.font.color.rgb = WHITE
    run.font.size = Pt(9.5)
    run.font.name = 'Calibri'

attacks = [
    ('Linkage / re-identification attack',
     'PPRL, Traditional De-id',
     'Joining two anonymized datasets can re-identify individuals (Netflix, AOL attacks). The core motivation for PETs.'),
    ('Membership inference attack',
     'FL, Synthetic Data',
     'An adversary queries a model to determine whether a specific student was in the training set.'),
    ('Gradient inversion / model inversion',
     'Federated Learning',
     'Gradients can be inverted to approximately reconstruct training records. Motivates DP-on-gradients; not mentioned anywhere in the FL module.'),
    ('Differencing attack',
     'DP, Perturbation',
     'Running two overlapping queries and subtracting reveals individual records. The classic DP origin story — mentioned in limitations text but not simulated.'),
    ('Homogeneity attack on k-anonymity',
     'Traditional (k-Anon)',
     'When all k records share the same sensitive attribute, grouping provides zero protection. Listed in limitations but not demonstrated interactively.'),
    ('Background knowledge attack',
     'Traditional (k-Anon)',
     'Listed in limitations text only. An adversary with auxiliary information can narrow a k-group to a single individual.'),
]
for ri, (attack, modules, why) in enumerate(attacks):
    fill = alt_fill[ri % 2]
    for ci, text in enumerate([attack, modules, why]):
        cell = tbl4.cell(ri + 1, ci)
        shade_cell(cell, fill)
        thin_border(cell)
        run = cell.paragraphs[0].add_run(text)
        run.font.size = Pt(9)
        run.font.name = 'Calibri'
        run.font.color.rgb = GREY_TEXT

spacer()

h2('4.2  PET Concepts Not Covered')

tbl5 = doc.add_table(rows=9, cols=2)
tbl5.style = 'Table Grid'
tbl5.alignment = WD_TABLE_ALIGNMENT.LEFT

for ci, h in enumerate(['Missing Concept', 'Gap Description']):
    cell = tbl5.cell(0, ci)
    shade_cell(cell, DARK_BLUE)
    thin_border(cell)
    run = cell.paragraphs[0].add_run(h)
    run.bold = True
    run.font.color.rgb = WHITE
    run.font.size = Pt(9.5)
    run.font.name = 'Calibri'

concepts = [
    ('Privacy budget composition (sequential & parallel)',
     'Running DP queries multiple times depletes the budget multiplicatively. Not shown or tracked in the UI. Critical for any real deployment.'),
    ('Approximate DP — the (ε, δ) parameter',
     'The δ parameter (probability of a catastrophic privacy failure) is entirely absent. All examples show pure ε-DP only.'),
    ('Secure Aggregation (in FL)',
     'Used to prevent the coordinator from seeing individual institution gradients. Separate from MPC but closely related; not mentioned in the FL module.'),
    ('l-diversity and t-closeness',
     'Extensions to k-anonymity that address homogeneity and distribution attacks. Mentioned once in the k-anonymity description but not simulated.'),
    ('Synthetic data memorization risk',
     'A poorly trained generative model can memorize outlier records. The simulation shows synthetic data as perfectly safe, which is not accurate for all generation methods.'),
    ('Zero-Knowledge Proofs (ZKPs)',
     'Increasingly used in credential, audit, and regulatory contexts. Not mentioned as a PET category.'),
    ('Privacy-Utility Pareto frontier',
     'No cross-module visualization showing which technique is optimal for a given utility requirement and threat model.'),
    ('Output type column — all 8 modules show *TBD*',
     'Whether output is aggregate-only vs. individual-record-level is a critical FERPA consideration. The column is promised in every tradeoff row but never filled in.'),
]
for ri, (concept, gap) in enumerate(concepts):
    fill = alt_fill[ri % 2]
    for ci, text in enumerate([concept, gap]):
        cell = tbl5.cell(ri + 1, ci)
        shade_cell(cell, fill)
        thin_border(cell)
        run = cell.paragraphs[0].add_run(text)
        run.font.size = Pt(9)
        run.font.name = 'Calibri'
        run.font.color.rgb = GREY_TEXT

spacer()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION 5 — PRIORITY REMEDIATION LIST
# ══════════════════════════════════════════════════════════════════════════════

h1('5  Priority Remediation List')

body('Ordered by impact. No code changes have been made; this list is for planning purposes.')

priorities = [
    ('P1 — Fix the LDP correction formula',
     'app.js:673',
     'Change the denominator from (2 * p - 1) to p. Also reconcile the formula comment in content.js:309. This is an outright bug affecting scientific accuracy.',
     CORAL),
    ('P2 — Fill in the *TBD* Output Type rows',
     'content.js — all 8 module tradeoff arrays',
     'Every module shows [\'Output type\', \'*TBD*\', \'val-med\']. This visible placeholder undermines credibility with the target audience.',
     AMBER),
    ('P3 — Add a differencing attack demo to the DP module',
     'app.js — renderDP() / renderGlobalTab()',
     'Show two overlapping queries whose difference reveals an individual count. This is the canonical DP motivation and the highest-value educational addition.',
     BLUE_TEXT),
    ('P4 — Add a cumulative ε budget meter to the DP module',
     'app.js — renderGlobalTab(), dpRunQuery()',
     'Track total ε spent across query runs. Display a budget bar (e.g., "Spent: 3.0 / 5.0 ε"). Teaches sequential composition.',
     BLUE_TEXT),
    ('P5 — Fix the PPRL matching to use hashes',
     'app.js:814 — pprlLink()',
     'Replace the id-shortcut match with hash-based comparison. Add one deliberate false positive to demonstrate Bloom filter probabilistic matching.',
     BLUE_TEXT),
    ('P6 — Remove or hide "True rate" from Local DP results view',
     'app.js:688 — ldpRunSurvey()',
     'The true rate should not be visible to the "collector" side of the simulation. Move it to an analyst/researcher view or remove it entirely.',
     GREY_TEXT),
]

for pri, (title, location, description, color) in enumerate(priorities):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    r1 = p.add_run(title)
    r1.bold = True
    r1.font.size  = Pt(10.5)
    r1.font.color.rgb = color
    r1.font.name  = 'Calibri'

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after  = Pt(1)
    p2.paragraph_format.left_indent  = Inches(0.2)
    r_loc = p2.add_run(f'File: {location}')
    r_loc.font.size  = Pt(9)
    r_loc.font.name  = 'Courier New'
    r_loc.font.color.rgb = RGBColor(0x55, 0x66, 0x88)

    body(description, left_indent=True)


# ══════════════════════════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════════════════════════

out_path = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                        'PET Simulation Review.docx')
doc.save(out_path)
print(f'Saved: {out_path}')
